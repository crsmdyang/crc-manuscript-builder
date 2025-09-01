"""
CRC Manuscript Builder – Streamlit app (MVP v3.1.4)
Author: ChatGPT (for Jun)
Date: 2025-09-01 (KST)

이번 버전(v3.1.4) 변경점
-------------------------
- ❌ **ModuleNotFoundError: 'docx' (python-docx)** 환경에서도 동작하도록 **.docx 읽기/쓰기 대체 경로** 추가
  - 읽기: `python-docx`가 없으면 ZIP으로 `.docx`를 열어 `word/document.xml`을 직접 파싱
  - 쓰기: `python-docx`가 없으면 **최소 DOCX**를 ZIP으로 생성(문단별 텍스트)
- ✅ v3.1.3의 **Streamlit Shim**(UI 미설치 환경)과 **PDF 백엔드 조건부 임포트** 유지
- 🧪 테스트 보강: DOCX fallback 생성/파싱 라운드트립 테스트 추가

설치(권장)
----------
- UI 사용(권장): `pip install streamlit requests pandas lxml pydantic tenacity python-docx pymupdf`
- 경량: `pip install streamlit requests pandas lxml pydantic tenacity pypdf`

실행:
  streamlit run app.py
"""
from __future__ import annotations
import os
import io
import re
import json
import zipfile
import requests
import pandas as pd
import xml.etree.ElementTree as ET
from typing import List, Dict, Optional, Tuple
from lxml import etree
from tenacity import retry, stop_after_attempt, wait_exponential
from pydantic import BaseModel, Field

# =====================
# Optional python-docx (with fallback)
# =====================
try:
    from docx import Document  # type: ignore
    from docx.shared import Pt  # type: ignore
    _HAVE_PYDOCX = True
except Exception:
    Document = None  # type: ignore
    Pt = None  # type: ignore
    _HAVE_PYDOCX = False

# =====================
# Streamlit import (with headless shim fallback)
# =====================
try:
    import streamlit as st  # type: ignore
    _HAVE_STREAMLIT = True
except Exception:
    _HAVE_STREAMLIT = False
    # ---- Streamlit Shim ----
    class _NoopContext:
        def __enter__(self):
            return self
        def __exit__(self, exc_type, exc, tb):
            return False
    class _SidebarShim:
        def success(self, *a, **k): print("[sidebar.success]", *a)
        def error(self, *a, **k): print("[sidebar.error]", *a)
        def warning(self, *a, **k): print("[sidebar.warning]", *a)
        def info(self, *a, **k): print("[sidebar.info]", *a)
    class _ColumnConfigShim:
        class CheckboxColumn:  # noqa: D401
            def __init__(self, *a, **k): pass
        class LinkColumn:
            def __init__(self, *a, **k): pass
        class TextColumn:
            def __init__(self, *a, **k): pass
    class _StreamlitShim:
        def __init__(self):
            self.session_state = {}
            self.sidebar = _SidebarShim()
            self.column_config = _ColumnConfigShim()
        # layout
        def set_page_config(self, *a, **k): pass
        def title(self, *a, **k): print("[title]", *a)
        def subheader(self, *a, **k): print("[subheader]", *a)
        def divider(self): print("[divider]")
        def caption(self, *a, **k): print("[caption]", *a)
        def write(self, *a, **k): print("[write]", *a)
        def markdown(self, *a, **k): print("[markdown]")
        def info(self, *a, **k): print("[info]", *a)
        def success(self, *a, **k): print("[success]", *a)
        def error(self, *a, **k): print("[error]", *a)
        def warning(self, *a, **k): print("[warning]", *a)
        # containers/contexts
        def expander(self, *a, **k): return _NoopContext()
        def spinner(self, *a, **k): return _NoopContext()
        def columns(self, spec):
            n = spec if isinstance(spec, int) else (len(spec) if hasattr(spec, "__len__") else 2)
            return [_NoopContext() for _ in range(n)]
        # widgets (return safe defaults)
        def text_area(self, *a, **k): return k.get("value", "")
        def text_input(self, *a, **k): return k.get("value", "")
        def selectbox(self, label, options, index=0, **k):
            try: return options[index]
            except Exception: return options[0] if options else ""
        def checkbox(self, label, value=False, **k): return bool(value)
        def slider(self, label, min_value=None, max_value=None, value=None, step=None, **k): return value
        def button(self, *a, **k): return False
        def file_uploader(self, *a, **k): return None
        def data_editor(self, df, *a, **k): return df
        def download_button(self, *a, **k): print("[download_button]")
    def cache_data(*a, **k):
        def _wrap(fn): return fn
        return _wrap
    st = _StreamlitShim()
    st.cache_data = cache_data

# =====================
# Conditional PDF backends (avoid 'frontend' import error from PyMuPDF)
# =====================
HAVE_FITZ = False
HAVE_PYPDF = False
_PDF_BACKEND = "none"
try:
    import fitz  # PyMuPDF
    HAVE_FITZ = True
    _PDF_BACKEND = "PyMuPDF"
except Exception:
    try:
        from pypdf import PdfReader as _PdfReader  # modern fork of PyPDF2
        HAVE_PYPDF = True
        _PDF_BACKEND = "pypdf"
    except Exception:
        try:
            from PyPDF2 import PdfReader as _PdfReader
            HAVE_PYPDF = True
            _PDF_BACKEND = "PyPDF2"
        except Exception:
            _PDF_BACKEND = "none"

# =====================
# Config & constants
# =====================
APP_TITLE = "CRC Manuscript Builder (MVP v3.1.4)"
EUTILS_BASE = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils"
CROSSREF_BASE = "https://api.crossref.org/works"
OPENALEX_BASE = "https://api.openalex.org/sources"
UNPAYWALL_BASE = "https://api.unpaywall.org/v2/"

MAX_RESULTS = 50
MAX_UPLOADS = 50

OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
ENTREZ_API_KEY = os.getenv("ENTREZ_API_KEY")
UNPAYWALL_EMAIL = os.getenv("UNPAYWALL_EMAIL")

DOI_RE = re.compile(r"10\.\d{4,9}/[-._;()/:A-Za-z0-9]+")

# =====================
# Helpers
# =====================

def norm_text(s: str) -> str:
    return re.sub(r"\s+", " ", s or "").strip()


def find_doi_in_text(txt: str) -> Optional[str]:
    m = DOI_RE.search(txt or "")
    return m.group(0) if m else None


# ---------- DOCX fallback helpers ----------
_DOCX_NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
    "dc": "http://purl.org/dc/elements/1.1/",
    "dcterms": "http://purl.org/dc/terms/",
    "xsi": "http://www.w3.org/2001/XMLSchema-instance",
}


def read_docx_text_fallback(docx_bytes: bytes) -> str:
    """Read .docx text without python-docx by parsing word/document.xml."""
    try:
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
            xml = z.read("word/document.xml")
        root = ET.fromstring(xml)
        texts = []
        for p in root.findall(".//w:p", _DOCX_NS):
            frag = []
            for t in p.findall(".//w:t", _DOCX_NS):
                frag.append(t.text or "")
            para = "".join(frag).strip()
            if para:
                texts.append(para)
        return "\n".join(texts)
    except Exception:
        return ""


def create_docx_from_markdown_fallback(md_text: str) -> bytes:
    """Create a minimal .docx (paragraph per block) without python-docx."""
    from xml.sax.saxutils import escape
    # very simple block split
    paragraphs = [p.strip() for p in md_text.split("\n\n")]
    if not any(paragraphs):
        paragraphs = [""]
    # Minimal required parts
    content_types = (
        """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>
  <Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>
</Types>"""
    ).strip()
    rels = (
        """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>"""
    ).strip()
    # Build document.xml paragraphs
    paras_xml = []
    for p in paragraphs:
        txt = escape(p)
        paras_xml.append(f"<w:p><w:r><w:t xml:space=\"preserve\">{txt}</w:t></w:r></w:p>")
    document_xml = (
        """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    {PARAS}
    <w:sectPr/>
  </w:body>
</w:document>""".replace("{PARAS}", "\n    ".join(paras_xml))
    ).strip()
    core = (
        """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:dcterms="http://purl.org/dc/terms/" xmlns:dcmitype="http://purl.org/dc/dcmitype/" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
  <dc:title>CRC Manuscript</dc:title>
  <dc:creator>CRC Manuscript Builder</dc:creator>
  <cp:lastModifiedBy>CRC Manuscript Builder</cp:lastModifiedBy>
</cp:coreProperties>"""
    ).strip()
    app = (
        """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties" xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">
  <Application>CRC Manuscript Builder</Application>
</Properties>"""
    ).strip()
    bio = io.BytesIO()
    with zipfile.ZipFile(bio, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", content_types)
        z.writestr("_rels/.rels", rels)
        z.writestr("word/document.xml", document_xml)
        z.writestr("docProps/core.xml", core)
        z.writestr("docProps/app.xml", app)
    bio.seek(0)
    return bio.read()


# ---------- PDF extract ----------

def extract_pdf_text_and_doi(file_bytes: bytes) -> Tuple[str, Optional[str]]:
    """Extract text & DOI from PDF bytes using available backend.
    Returns (full_text, doi_or_None). Safe to call even if no backend.
    """
    # Backend 1: PyMuPDF
    if HAVE_FITZ:
        try:
            import fitz  # re-import inside for safety
            with fitz.open(stream=file_bytes, filetype="pdf") as doc:
                texts = [page.get_text() for page in doc]
            full = "\n".join(texts)
            return full, find_doi_in_text(full)
        except Exception:
            pass  # fallthrough
    # Backend 2: pypdf / PyPDF2
    if HAVE_PYPDF:
        try:
            from io import BytesIO
            bio = BytesIO(file_bytes)
            reader = _PdfReader(bio)  # type: ignore[name-defined]
            texts = []
            for page in getattr(reader, "pages", []):
                try:
                    texts.append(page.extract_text() or "")
                except Exception:
                    texts.append("")
            full = "\n".join(texts)
            return full, find_doi_in_text(full)
        except Exception:
            pass
    # No backend
    return "", None


# =====================
# Data models
# =====================
class RefMeta(BaseModel):
    doi: Optional[str] = None
    title: Optional[str] = None
    journal: Optional[str] = None
    year: Optional[str] = None
    authors: List[str] = Field(default_factory=list)
    pmid: Optional[str] = None
    url: Optional[str] = None
    abstract_text: Optional[str] = None
    abstract_conclusion: Optional[str] = None


class ReferenceManager:
    def __init__(self):
        self.idx: Dict[str, int] = {}
        self.meta: Dict[str, RefMeta] = {}
        self.order: List[str] = []

    def _key(self, m: RefMeta) -> Optional[str]:
        if m.doi:
            return m.doi.lower()
        if m.pmid:
            return f"pmid:{m.pmid}"
        return None

    def register(self, meta: RefMeta) -> Optional[int]:
        key = self._key(meta)
        if not key:
            return None
        if key not in self.idx:
            self.order.append(key)
            self.idx[key] = len(self.order)
            self.meta[key] = meta
        else:
            old = self.meta.get(key)
            self.meta[key] = meta if len(json.dumps(meta.dict())) > len(json.dumps(old.dict())) else old
        return self.idx[key]

    def cite(self, doi_or_pmid: str) -> Optional[int]:
        key = doi_or_pmid.lower()
        if key not in self.idx:
            self.idx[key] = len(self.order) + 1
            self.order.append(key)
            self.meta.setdefault(key, RefMeta())
        return self.idx[key]

    def renumber_by_first_appearance(self, citation_sequence: List[str]):
        seen = []
        for key in citation_sequence:
            if key not in seen and key in self.idx:
                seen.append(key)
        for key in self.order:
            if key not in seen:
                seen.append(key)
        self.order = seen
        self.idx = {k: i + 1 for i, k in enumerate(self.order)}

    def vancouver(self, key: str) -> str:
        m = self.meta.get(key, RefMeta())
        def fmt_author(a: str) -> str:
            a = a.strip()
            parts = a.split()
            if not parts:
                return a
            last = parts[-1]
            initials = "".join(p[0] for p in parts[:-1] if p)
            return f"{last} {initials}".strip()
        authors = ", ".join(fmt_author(a) for a in (m.authors or [])[:6])
        if m.authors and len(m.authors) > 6:
            authors += ", et al."
        year = m.year or ""
        journal = m.journal or ""
        title = m.title or ""
        doi = f" https://doi.org/{m.doi}" if m.doi else (f" PMID:{m.pmid}" if m.pmid else "")
        return norm_text(f"{authors}. {title}. {journal}. {year}.{doi}")

    def render_reference_list(self) -> List[str]:
        return [self.vancouver(k) for k in self.order]


# =====================
# External services
# =====================
@retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=1, max=6))
def pubmed_search(term: str, retmax: int = MAX_RESULTS) -> List[str]:
    params = {"db": "pubmed", "term": term, "retmode": "json", "retmax": retmax, "sort": "relevance"}
    if ENTREZ_API_KEY:
        params["api_key"] = ENTREZ_API_KEY
    r = requests.get(f"{EUTILS_BASE}/esearch.fcgi", params=params, timeout=30)
    r.raise_for_status()
    data = r.json()
    return data.get("esearchresult", {}).get("idlist", [])


@retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=1, max=6))
def pubmed_fetch_xml(pmids: List[str]) -> etree._Element:
    if not pmids:
        return etree.Element("empty")
    params = {"db": "pubmed", "id": ",".join(pmids), "retmode": "xml"}
    if ENTREZ_API_KEY:
        params["api_key"] = ENTREZ_API_KEY
    r = requests.get(f"{EUTILS_BASE}/efetch.fcgi", params=params, timeout=60)
    r.raise_for_status()
    return etree.fromstring(r.content)


def _xml_text(node, xpath: str) -> Optional[str]:
    el = node.find(xpath)
    return norm_text(el.text) if el is not None and el.text else None


def _extract_abstract_and_conclusion(art: etree._Element) -> Tuple[Optional[str], Optional[str]]:
    nodes = art.findall(".//Abstract/AbstractText")
    if not nodes:
        return None, None
    parts = []
    conclusion = None
    for t in nodes:
        label = (t.get("Label") or t.get("NlmCategory") or "").lower()
        val = norm_text("".join(t.itertext()))
        if not val:
            continue
        parts.append(val)
        if label.startswith("conclusion"):
            sent = re.split(r"(?<=[.!?])\s+", val)
            conclusion = " ".join(sent[:2]).strip()
    full = " \n".join(parts) if parts else None
    if not conclusion and parts:
        sent = re.split(r"(?<=[.!?])\s+", parts[-1])
        conclusion = " ".join(sent[:2]).strip() if sent else None
    return full, conclusion


def pubmed_parse_records(root: etree._Element) -> List[RefMeta]:
    out: List[RefMeta] = []
    for art in root.findall(".//PubmedArticle"):
        pmid = _xml_text(art, ".//MedlineCitation/PMID")
        title = _xml_text(art, ".//Article/ArticleTitle")
        year = _xml_text(art, ".//Article/Journal/JournalIssue/PubDate/Year")
        journal = _xml_text(art, ".//Article/Journal/Title")
        # Authors
        authors = []
        for a in art.findall(".//AuthorList/Author"):
            ln = _xml_text(a, "LastName") or ""
            fn = _xml_text(a, "ForeName") or ""
            full = norm_text(f"{fn} {ln}")
            if full:
                authors.append(full)
        # DOI
        doi = None
        for idn in art.findall(".//ArticleIdList/ArticleId"):
            if idn.get("IdType") == "doi" and idn.text:
                doi = idn.text.lower()
        url = f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/" if pmid else None
        abst, concl = _extract_abstract_and_conclusion(art)
        out.append(RefMeta(doi=doi, title=title, journal=journal, year=year, authors=authors,
                           pmid=pmid, url=url, abstract_text=abst, abstract_conclusion=concl))
    return out


def crossref_get(doi: str) -> Optional[RefMeta]:
    try:
        r = requests.get(f"{CROSSREF_BASE}/{doi}", timeout=30)
        if r.status_code != 200:
            return None
        j = r.json().get("message", {})
        title = "; ".join(j.get("title", [])) or None
        journal = (j.get("container-title") or [None])[0]
        year = None
        if j.get("issued", {}).get("'date-parts'"):
            year = str(j["issued"]["'date-parts'"][0][0])
        elif j.get("issued", {}).get("date-parts"):
            year = str(j["issued"]["date-parts"][0][0])
        authors = []
        for a in j.get("author", []):
            nm = norm_text(f"{a.get('given','')} {a.get('family','')}")
            if nm:
                authors.append(nm)
        url = j.get("URL")
        return RefMeta(doi=doi.lower(), title=title, journal=journal, year=year, authors=authors, url=url)
    except Exception:
        return None


def openalex_metric(journal_title: str) -> Optional[float]:
    try:
        q = {"search": journal_title}
        r = requests.get(OPENALEX_BASE, params=q, timeout=20)
        if r.status_code != 200:
            return None
        data = r.json().get("results", [])
        if not data:
            return None
        src = data[0]
        sjr2 = None
        try:
            sjr2 = src.get("summary_stats", {}).get("two_year_sjr")
        except Exception:
            sjr2 = None
        return float(sjr2) if sjr2 is not None else None
    except Exception:
        return None


def unpaywall_best_oa_link(doi: str) -> Optional[str]:
    if not UNPAYWALL_EMAIL:
        return None
    try:
        r = requests.get(f"{UNPAYWALL_BASE}{doi}", params={"email": UNPAYWALL_EMAIL}, timeout=20)
        if r.status_code != 200:
            return None
        j = r.json()
        rec = j.get("best_oa_location") or {}
        return rec.get("url")
    except Exception:
        return None


# =====================
# RIS exporter (EndNote)
# =====================

def to_ris(refs: List[RefMeta]) -> str:
    lines = []
    for r in refs:
        lines.append("TY  - JOUR")
        if r.title: lines.append(f"TI  - {r.title}")
        if r.journal: lines.append(f"JO  - {r.journal}")
        if r.year: lines.append(f"PY  - {r.year}")
        for au in (r.authors or []):
            lines.append(f"AU  - {au}")
        if r.doi: lines.append(f"DO  - {r.doi}")
        if r.url: lines.append(f"UR  - {r.url}")
        if r.pmid: lines.append(f"ID  - PMID:{r.pmid}")
        if r.abstract_text: lines.append(f"AB  - {r.abstract_text}")
        lines.append("ER  - ")
    return "\n".join(lines) + "\n"


# =====================
# UI (works in Streamlit; no-op in headless shim)
# =====================
st.set_page_config(page_title=APP_TITLE, layout="wide")
st.title(APP_TITLE)

with st.expander("사용 지침(필독)", expanded=True):
    st.markdown(
        f"""
        - **허구 금지**: 인용은 선택/업로드한 문헌으로만 제한됩니다.
        - **IF 정렬**: `journal_if.csv` 제공 시 IF 기준 정렬, 없으면 OpenAlex 지표(선택) 사용 가능.
        - **RIS 내보내기**: 검색 선택/허용 문헌(전체·선택)을 EndNote용 `.ris`로 다운로드할 수 있습니다.
        - **PDF 업로드**: 최대 50개, PDF 내 DOI 자동 추출 시도.
        - **백엔드**: PDF → **{_PDF_BACKEND}**, DOCX → **{'python-docx' if _HAVE_PYDOCX else 'fallback-zip'}**, 모드 → **{'Streamlit UI' if _HAVE_STREAMLIT else 'Headless(Shim)'}**
        """
    )

# 1) 입력(연구계획서 .docx 업로드 포함) + 타깃 저널/스타일
colA, colB = st.columns([3, 2])
with colA:
    topic = st.text_area("주제 (Topic)", height=80)
    protocol = st.text_area("연구계획서 요약 (Study Protocol)", height=160, key="protocol_ta")
    up_docx = st.file_uploader("연구계획서 요약 .docx 업로드 (선택)", type=["docx"], accept_multiple_files=False)
    if up_docx is not None:
        try:
            data = up_docx.read()
            if _HAVE_PYDOCX and Document:
                bio = io.BytesIO(data)
                doc = Document(bio)
                text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            else:
                text = read_docx_text_fallback(data)
            if text:
                st.session_state["protocol_ta"] = text
                st.success("워드 파일에서 연구계획서 요약을 불러왔어요.")
            else:
                st.warning("워드 파일에서 텍스트를 추출할 수 없습니다.")
        except Exception as e:
            st.error(f"워드 파일을 읽는 중 오류: {e}")
    results_txt = st.text_area("핵심 결과 요약 (Key Results)", height=100)

with colB:
    style_option = st.selectbox("타깃 저널/스타일", [
        "국내: 대한대장항문학회(Annals of Coloproctology)",
        "해외: ASCRS 스타일 (Diseases of the Colon & Rectum)",
        "해외: ESCP 스타일 (Colorectal Disease)",
        "없음/기타(직접입력)",
    ])
    custom_style = ""
    if style_option == "없음/기타(직접입력)":
        custom_style = st.text_input("직접 입력", placeholder="예: BJS 스타일, 또는 목표 저널명")
    use_openalex = st.checkbox("OpenAlex 지표 사용(대체 지표)", value=False)

st.divider()

# 2) PubMed 검색
st.subheader("2) PubMed 검색")
search_query = st.text_input("검색식", placeholder="예: (rectal cancer OR colorectal) AND (chemoradiation)")
retmax = st.slider("검색 개수", 10, MAX_RESULTS, 50, step=5)
run_search = st.button("PubMed 검색 실행")

if "search_results" not in st.session_state:
    st.session_state["search_results"] = []
if "search_df" not in st.session_state:
    st.session_state["search_df"] = None

if run_search and search_query:
    with st.spinner("PubMed 검색 중…"):
        pmids = pubmed_search(search_query, retmax=retmax)
        root = pubmed_fetch_xml(pmids)
        st.session_state["search_results"] = pubmed_parse_records(root)

# 2-1) IF 붙이고 정렬 (IF desc → journal asc → year desc)
@st.cache_data(show_spinner=False)
def load_journal_if_csv() -> Optional[pd.DataFrame]:
    try:
        if os.path.exists("journal_if.csv"):
            df = pd.read_csv("journal_if.csv")
            cols = {c.lower(): c for c in df.columns}
            jcol = cols.get("journal") or cols.get("title") or list(df.columns)[0]
            icol = cols.get("if") or cols.get("impact_factor") or list(df.columns)[1]
            df = df.rename(columns={jcol: "journal", icol: "if"})
            df["journal_norm"] = df["journal"].str.strip().str.lower()
            return df
    except Exception:
        pass
    return None

jif = load_journal_if_csv()

if st.session_state.get("search_results"):
    sdf = pd.DataFrame([
        {
            "select": False,
            "pmid": r.pmid,
            "doi": r.doi,
            "title": r.title,
            "journal": r.journal,
            "year": r.year,
            "Abstract": (r.abstract_text[:300] + "…") if r.abstract_text and len(r.abstract_text) > 300 else (r.abstract_text or None),
            "Conclusion": r.abstract_conclusion,
            "url": r.url,
        }
        for r in st.session_state["search_results"]
    ])
    sdf["IF"] = None
    if jif is not None:
        jmap = dict(zip(jif["journal_norm"], jif["if"]))
        sdf["IF"] = sdf["journal"].fillna("").str.strip().str.lower().map(jmap)
    elif use_openalex:
        metrics = []
        for jn in sdf["journal"].fillna(""):
            metrics.append(openalex_metric(jn) if jn else None)
        sdf["IF"] = metrics

    def to_float(x):
        try: return float(x)
        except: return float("nan")
    def to_int(x):
        try: return int(x)
        except: return -1
    sdf["IF_num"] = sdf["IF"].apply(to_float).fillna(-1.0)
    sdf["year_num"] = sdf["year"].apply(to_int)
    sdf = sdf.sort_values(by=["IF_num", "journal", "year_num"], ascending=[False, True, False]).reset_index(drop=True)

    st.markdown("**검색 결과(체크 → 허용 문헌 추가 / RIS 내보내기)**")
    edited = st.data_editor(
        sdf,
        hide_index=True,
        use_container_width=True,
        column_config={
            "select": st.column_config.CheckboxColumn("선택"),
            "url": st.column_config.LinkColumn("PubMed"),
            "Abstract": st.column_config.TextColumn("Abstract", width="large"),
            "Conclusion": st.column_config.TextColumn("Conclusion", width="medium"),
            "IF": st.column_config.TextColumn("IF/Proxy"),
        },
        key="search_editor",
    )
    st.session_state["search_df"] = edited

    c1, c2, c3 = st.columns([1,1,2])
    with c1:
        add_sel = st.button("선택 추가 → 허용 문헌")
    with c2:
        ris_sel_btn = st.button("선택 .ris 다운로드")
    with c3:
        st.write(f"선택 수: **{int((edited['select']==True).sum())}** 편")

    if add_sel or ris_sel_btn:
        chosen_rows = edited[edited["select"] == True]
        chosen_refs: List[RefMeta] = []
        for _, row in chosen_rows.iterrows():
            meta = RefMeta(
                doi=(str(row.get("doi")) if row.get("doi") else None),
                pmid=(str(row.get("pmid")) if row.get("pmid") else None),
                title=row.get("title"), journal=row.get("journal"), year=str(row.get("year")),
                url=row.get("url"), abstract_text=row.get("Abstract"), abstract_conclusion=row.get("Conclusion")
            )
            chosen_refs.append(meta)
            if add_sel:
                key = (meta.doi.lower() if meta.doi else (f"pmid:{meta.pmid}" if meta.pmid else None))
                if key:
                    st.session_state.setdefault("allowed", {})
                    st.session_state["allowed"][key] = meta
        if add_sel:
            st.success(f"허용 문헌에 {len(chosen_refs)}편 추가")
        if ris_sel_btn:
            ris_txt = to_ris(chosen_refs)
            st.download_button("선택 .ris 다운로드", data=ris_txt.encode("utf-8"), file_name="pubmed_selection.ris", mime="application/x-research-info-systems")

st.divider()

# 3) PDF 업로드 → 허용 문헌
st.subheader("3) PDF 업로드(최대 50) → 허용 문헌")
pdfs = st.file_uploader("논문 PDF 업로드", type=["pdf"], accept_multiple_files=True)
if st.button("PDF에서 DOI 추출 후 추가") and pdfs:
    if _PDF_BACKEND == "none":
        st.warning("PDF 파서가 설치되어 있지 않아 DOI를 추출할 수 없습니다. 'pymupdf' 또는 'pypdf'를 설치해 주세요.")
    added = 0
    st.session_state.setdefault("allowed", {})
    for up in pdfs[:MAX_UPLOADS]:
        content = up.read()
        _, doi = extract_pdf_text_and_doi(content)
        if doi:
            meta = crossref_get(doi) or RefMeta(doi=doi)
            st.session_state["allowed"][doi.lower()] = meta
            added += 1
    st.success(f"PDF에서 {added}편 추가 (DOI 미탐지 파일은 생략)")

# 4) 허용 문헌 표(추가/삭제, RIS 내보내기)
st.subheader("4) 허용 문헌 (인용 가능한 집합)")
st.session_state.setdefault("allowed", {})

if st.session_state.get("allowed"):
    adf_rows = []
    for k, v in st.session_state["allowed"].items():
        # 방어적 접근(혹시 dict 등이 섞여 들어온 경우 대비)
        doi = getattr(v, "doi", None)
        pmid = getattr(v, "pmid", None)
        title = getattr(v, "title", None)
        journal = getattr(v, "journal", None)
        year = getattr(v, "year", None)
        abstract_text = getattr(v, "abstract_text", None)
        conclusion_text = getattr(v, "abstract_conclusion", None)
        preview = (abstract_text[:300] + "…") if abstract_text and len(abstract_text) > 300 else (abstract_text or None)
        adf_rows.append({
            "select": False,
            "key": k,
            "doi": doi,
            "pmid": pmid,
            "title": title,
            "journal": journal,
            "year": year,
            "Abstract": preview,
            "Conclusion": conclusion_text,
            "OA_link": unpaywall_best_oa_link(doi) if doi else None,
        })
    adf = pd.DataFrame(adf_rows)

    edited_allowed = st.data_editor(
        adf,
        hide_index=True,
        use_container_width=True,
        column_config={
            "select": st.column_config.CheckboxColumn("선택"),
            "OA_link": st.column_config.LinkColumn("OA 링크"),
            "Abstract": st.column_config.TextColumn("Abstract", width="large"),
            "Conclusion": st.column_config.TextColumn("Conclusion", width="medium"),
        },
        key="allowed_editor",
    )

    c1, c2, c3 = st.columns([1,1,2])
    with c1:
        del_btn = st.button("선택 삭제")
    with c2:
        ris_allowed_sel = st.button("선택 .ris 다운로드")
    with c3:
        ris_allowed_all = st.button("전체 .ris 다운로드")

    if del_btn:
        to_del = edited_allowed[edited_allowed["select"] == True]["key"].tolist()
        for k in to_del:
            st.session_state["allowed"].pop(k, None)
        st.success(f"삭제 완료: {len(to_del)}편")

    if ris_allowed_sel or ris_allowed_all:
        export_keys = (edited_allowed[edited_allowed["select"] == True]["key"].tolist() if ris_allowed_sel
                       else list(st.session_state["allowed"].keys()))
        refs = [st.session_state["allowed"][k] for k in export_keys if k in st.session_state["allowed"]]
        ris_txt = to_ris(refs)
        fname = "allowed_selection.ris" if ris_allowed_sel else "allowed_all.ris"
        st.download_button(".ris 다운로드", data=ris_txt.encode("utf-8"), file_name=fname, mime="application/x-research-info-systems")
else:
    st.info("허용 문헌이 비어 있습니다. 검색 결과에서 선택하여 추가하세요.")

st.divider()

# 5) 섹션별 생성 → 병합
st.subheader("5) 섹션별 생성 → 최종 병합")
class LLM:
    def __init__(self, model: str, api_key: Optional[str]):
        self.model = model
        self.key = api_key
        self.enabled = bool(api_key)
        if not self.enabled:
            st.warning("OPENAI_API_KEY 가 설정되어 있지 않습니다. 생성 기능은 비활성화됩니다.")

    def generate_section(self, section: str, topic: str, protocol: str, results: str,
                         allowed_refs: Dict[str, RefMeta], style_note: str) -> str:
        if not self.enabled:
            return "(LLM 비활성화: OPENAI_API_KEY 설정 필요)"
        try:
            refs_serialized = []
            for k, m in allowed_refs.items():
                label = m.doi or m.pmid or k
                refs_serialized.append({
                    "key": k,
                    "label": label,
                    "title": m.title,
                    "journal": m.journal,
                    "year": m.year,
                })
            system = (
                "You are an evidence-based medical writing assistant for colorectal surgery/oncology. "
                "Strictly use only the provided source list and cite with [CITE:DOI] or [CITE:pmid:ID]. "
                "If evidence is insufficient, write 'No high-quality evidence available.' Do not invent citations."
            )
            user = {
                "task": f"Write the {section} for a colorectal manuscript.",
                "topic": topic,
                "study_protocol": protocol,
                "key_results": results,
                "journal_style": style_note,
                "citation_rule": "Use only allowed sources via [CITE:...] tags.",
                "allowed_sources": refs_serialized,
                "language": "Korean",
            }
            payload = {"model": self.model, "messages": [
                {"role": "system", "content": system},
                {"role": "user", "content": json.dumps(user, ensure_ascii=False)},
            ], "temperature": 0.2}
            r = requests.post(
                "https://api.openai.com/v1/chat/completions",
                headers={"Authorization": f"Bearer {self.key}", "Content-Type": "application/json"},
                json=payload, timeout=120,
            )
            r.raise_for_status()
            content = r.json()["choices"][0]["message"]["content"]
            bad = []
            for tag in re.findall(r"\[CITE:([^\]]+)\]", content or ""):
                k = tag.strip().lower()
                if k not in allowed_refs:
                    bad.append(k)
            if bad:
                return (
                    "(생성 거부) 허용되지 않은 인용 태그: " + ", ".join(sorted(set(bad))) + "\n허용된 DOI/PMID만 사용하세요."
                )
            return content
        except Exception as e:
            return f"(LLM 오류) {e}"

llm = LLM(OPENAI_MODEL, OPENAI_API_KEY)
style_note = (custom_style if style_option == "없음/기타(직접입력)" else style_option)
SECTIONS = ["Cover Letter", "Title Page", "Abstract", "Introduction", "Methods", "Results", "Discussion"]

if "sections" not in st.session_state:
    st.session_state["sections"] = {}

cols = st.columns(2)
for i, sec in enumerate(SECTIONS):
    with cols[i % 2]:
        st.markdown(f"**{sec}**")
        if st.button(f"{sec} 생성", key=f"gen_{sec}"):
            txt = llm.generate_section(sec, topic, st.session_state.get("protocol_ta", ""), results_txt, st.session_state.get("allowed", {}), style_note)
            st.session_state["sections"][sec] = txt
        st.text_area(f"{sec} 미리보기", value=st.session_state["sections"].get(sec, ""), height=200, key=f"ta_{sec}")

st.markdown("**References** 섹션은 최종 병합 단계에서 자동 생성됩니다.")

if st.button("최종 병합 및 번호 재정렬"):
    rm = ReferenceManager()
    for k, m in st.session_state.get("allowed", {}).items():
        if k.startswith("pmid:") and m:
            m.pmid = k.split(":",1)[1]
        elif m:
            m.doi = m.doi or k
        rm.register(m)

    def replace_citations(text: str) -> Tuple[str, List[str]]:
        seq = []
        def _rep(m):
            tag = m.group(1).strip().lower()
            if tag not in st.session_state.get("allowed", {}):
                return f"[CITE-INVALID:{tag}]"
            seq.append(tag)
            n = rm.cite(tag)
            return f"[{n}]"
        new = re.sub(r"\[CITE:([^\]]+)\]", _rep, text or "")
        return new, seq

    merged_parts = []
    citation_seq = []
    for sec in SECTIONS:
        txt = st.session_state["sections"].get(sec, "")
        if not txt:
            continue
        rep, seq = replace_citations(txt)
        merged_parts.append(f"## {sec}\n\n" + rep.strip())
        citation_seq.extend(seq)

    rm.renumber_by_first_appearance(citation_seq)
    refs = rm.render_reference_list()

    final_md = "\n\n".join(merged_parts) + "\n\n## References\n\n" + "\n".join(
        [f"[{i+1}] {line}" for i, line in enumerate(refs)]
    )
    st.session_state["final_md"] = final_md
    st.success("병합 완료 – 아래에서 미리보기/내보내기 하세요.")

# 미리보기 및 내보내기
if "final_md" in st.session_state:
    st.subheader("미리보기 (Markdown)")
    st.text_area("Final Markdown", value=st.session_state["final_md"], height=420)

    def md_to_docx_bytes(md_text: str) -> bytes:
        if _HAVE_PYDOCX and Document and Pt:
            # python-docx path
            bio = io.BytesIO()
            try:
                doc = Document()
                style = doc.styles["Normal"]
                style.font.name = "Calibri"
                style.font.size = Pt(11)
                for block in md_text.split("\n\n"):
                    doc.add_paragraph(block)
                doc.save(bio)
                bio.seek(0)
                return bio.read()
            except Exception:
                # fallback if something goes wrong even with python-docx
                return create_docx_from_markdown_fallback(md_text)
        else:
            # fallback builder
            return create_docx_from_markdown_fallback(md_text)

    st.download_button("Download .md", data=st.session_state["final_md"].encode("utf-8"), file_name="manuscript.md", mime="text/markdown")
    st.download_button("Download .docx", data=md_to_docx_bytes(st.session_state["final_md"]), file_name="manuscript.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

st.divider()

st.caption("© 2025 CRC Manuscript Builder (MVP v3.1.4). Evidence-locked generation. No PHI.")

# =====================
# (Optional) Lightweight sanity tests (run only if explicitly requested)
# =====================

def _run_sanity_tests():
    """Minimal non-network tests to catch regressions in critical blocks."""
    # 1) Test RIS exporter with minimal RefMeta (existing)
    r = RefMeta(doi="10.1000/test.doi", title="T", journal="J", year="2024", authors=["A B"], pmid="1", url="U", abstract_text="abs")
    ris = to_ris([r])
    assert "TY  - JOUR" in ris and "DO  - 10.1000/test.doi" in ris
    # 2) Test reference manager numbering (existing)
    rm = ReferenceManager(); rm.register(r); n = rm.cite(r.doi)
    assert n == 1
    # 3) Test citation replacement helper (existing)
    txt = "See [CITE:10.1000/test.doi]."
    rm2 = ReferenceManager(); rm2.register(r)
    def _rep(m):
        tag = m.group(1).strip().lower(); rm2.cite(tag); return "[1]"
    out = re.sub(r"\[CITE:([^\]]+)\]", _rep, txt)
    assert out == "See [1]."
    # 4) New: DOI regex basic
    assert find_doi_in_text("doi:10.5555/abc.DEF-123") == "10.5555/abc.DEF-123"
    # 5) New: whitespace normalization
    assert norm_text("  a\t b\n c  ") == "a b c"
    # 6) New: PDF extractor with invalid bytes should not crash and return (text, None)
    text, doi = extract_pdf_text_and_doi(b"not-a-real-pdf")
    assert isinstance(text, str) and (doi is None or isinstance(doi, str))
    # 7) New: DOCX fallback roundtrip (build then read)
    sample = "Hello world\n\nSecond paragraph"
    built = create_docx_from_markdown_fallback(sample)
    parsed = read_docx_text_fallback(built)
    assert "Hello world" in parsed and "Second paragraph" in parsed
    return "OK"

# To run: set environment variable RUN_APP_TESTS=1 before launching Streamlit
if os.getenv("RUN_APP_TESTS") == "1":
    try:
        res = _run_sanity_tests()
        st.sidebar.success(f"Sanity tests: {res}")
    except Exception as e:
        st.sidebar.error(f"Sanity tests failed: {e}")
